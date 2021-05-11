import os, docx
from datetime import datetime
from covmass.models import Zone, Update
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt

def updateDOCX():

  doc = docx.Document()
  doc.add_heading('Covid-19 Lagebericht', 0)

  p1 = doc.add_paragraph("")
  p1.paragraph_format.space_before = Pt(20)

  table = doc.add_table(rows=1, cols=6)
  table.style = "Light Grid Accent 1"
  table.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
  header_cells = table.rows[0].cells
  header_cells[0].text = "Ort"
  header_cells[1].text = "Bestätigte Infektionen"
  header_cells[2].text = "∆Infektionen"
  header_cells[3].text = "Todesfälle"
  header_cells[4].text = "∆Todesfälle"
  header_cells[5].text = "Quelle"

  header_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
  header_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
  header_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
  header_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
  header_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
  header_cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


  for zone in Zone.objects.all():
    cells = table.add_row().cells
    
    cells[0].text = zone.german_name
    cells[1].text = zone.infected.display_total()
    cells[2].text = zone.infected.display_new()
    cells[3].text = zone.deceased.display_total()
    cells[4].text = zone.deceased.display_new()
    cells[5].text = zone.source.name
    
    cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    cells[1].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    cells[2].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    cells[3].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    cells[4].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    
    p2 = doc.add_paragraph(f"Diese Zahlen wurden von NEXUS ETH Zürich am {str(datetime.today().day).zfill(2)}.{str(datetime.today().month).zfill(2)}.{datetime.today().year} um 15:30")
  # p2 = doc.add_paragraph(f"Diese Zahlen wurden von NEXUS ETH Zürich am {str(datetime.today().day).zfill(2)}.{str(datetime.today().month).zfill(2)}.{datetime.today().year} um {str(datetime.today().hour).zfill(2)}:{str(datetime.today().minute).zfill(2)} aktualisiert.")
  p2.paragraph_format.space_before = Pt(20)

  doc.save("static_cdn/media_root/lagebericht.docx")
  print("DOCX is saved")
  
  return "Updated DOCX file"