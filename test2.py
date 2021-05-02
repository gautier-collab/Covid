import docx
import os

doc = docx.Document()
doc.add_heading('Covid-19 Lagebericht', 0)

table = doc.add_table(rows=1, cols=6)
table.style = "Medium Shading 1 Accent 1"
header_cells = table.rows[0].cells
header_cells[0].text = "Ort"
header_cells[1].text = "Bestätigte Infektionen"
header_cells[2].text = "∆Infektionen"
header_cells[3].text = "Todesfälle"
header_cells[4].text = "∆Todesfälle"
header_cells[5].text = "Quelle"

for i in range(5):
  row_cells = table.add_row().cells
  row_cells[0].text = "oklm"
  row_cells[1].text = "oklm"
  row_cells[2].text = "oklm"
  row_cells[3].text = "oklm"
  row_cells[4].text = "oklm"
  row_cells[5].text = "oklm"

doc.save("test.docx")

print("doc is saved")
