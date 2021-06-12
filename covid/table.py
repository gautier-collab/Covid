import os, docx
from datetime import datetime
from covid.models import Zone, Update, Update_zh
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from django.conf import settings
from os import walk

# function creating a new Word document from the newly fetched values
def updateDOCX():

  doc = docx.Document()
  doc.add_heading('Covid-19 Lagebericht', 0)

  p1 = doc.add_paragraph("")
  p1.paragraph_format.space_before = Pt(20)

  # header row content
  table = doc.add_table(rows=1, cols=6)
  table.style = "Light Grid Accent 1"
  table.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
  header_cells = table.rows[0].cells
  header_cells[0].text = "Ort"
  header_cells[1].text = "Bestätigte Infektionen"
  header_cells[2].text = "∆Infektionen"
  header_cells[3].text = "Todesfälle"
  header_cells[4].text = "∆Todesfälle"
  header_cells[5].text = "Quelle"

  # center vertically the values contained in the header cells
  header_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
  header_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
  header_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
  header_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
  header_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
  header_cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

  # iterate over every Zone instance of the database to define a new row
  for zone in Zone.objects.all():
    cells = table.add_row().cells
    
    # row content
    cells[0].text = zone.german_name
    cells[1].text = zone.infected.display_total()
    cells[2].text = zone.infected.display_new()
    cells[3].text = zone.deceased.display_total()
    cells[4].text = zone.deceased.display_new()
    cells[5].text = zone.source.name
    
    # center vertically the values contained in the cells
    cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # align to the right the numerical values
    cells[1].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    cells[2].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    cells[3].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    cells[4].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    
  # Include a sentence giving the time of the most recent update
  p2 = doc.add_paragraph(f"Diese Zahlen wurden von NEXUS ETH Zürich am {Update.objects.all().last().time} aktualisiert.")
  p2.paragraph_format.space_before = Pt(20)

  # get the last update time of Zürich values and include a sentence if the 'display' attribute is True
  update_zh = Update_zh.objects.all().last()
  if update_zh.display :
    p3 = doc.add_paragraph(f"Zürich letzte Aktualisierung: {update_zh.time}")
    p3.paragraph_format.space_before = Pt(10)

  # delete the already existing Word document
  path=f"{settings.BASE_DIR}/static_cdn/media_root"
  for (dirpath, dirnames, filenames) in walk(path):
    for filename in filenames:
      if "lagebericht" in filename:
        os.remove(f"{settings.BASE_DIR}/static_cdn/media_root/{filename}")
        
  # save the newly defined Word document
  doc.save("static_cdn/media_root/lagebericht.docx")
  print("DOCX is saved")
  
  return "Word document is replaced"